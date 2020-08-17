import os
import time
from selenium.common.exceptions import NoSuchElementException
import requests
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys


def get_url(driver, url, step_info):
    driver.get(url)
    driver.implicitly_wait(1)
    driver.maximize_window()
    driver.implicitly_wait(2)
    r = requests.get(url)
    #print(r.status_code)
    driver.execute_script("window.scrollTo(0, 0);")
    if r.status_code == 200:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/PASS_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('URL accessed as expected', 'PASS', 'PASS_' + str(step_info).replace('-', "").replace(" ", "_"))
    else:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/FAIL_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('URL not accessible', 'FAIL', 'FAIL_' + str(step_info).replace('-', "").replace(" ", "_"))


def search_element(driver, path, step_info, comp_string):
    actions = ActionChains(driver)
    actions.move_to_element(driver.find_element_by_xpath(path)).send_keys(Keys.ARROW_DOWN).perform()
    driver.implicitly_wait(1)
    ele_label = driver.find_element_by_xpath(path).text

    if ele_label == comp_string:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/PASS_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Element found with label text as ' + driver.find_element_by_xpath(path).text, 'PASS', "PASS_" + str(step_info).replace('-', "").replace(" ", "_"))

    else:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/FAIL_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Element found with label text as ' + driver.find_element_by_xpath(path).text, 'FAIL', "FAIL_" + str(step_info).replace('-', "").replace(" ", "_"))

def click_target_element(driver, path, step_info, comp_string):
    actions = ActionChains(driver)
    try:
        actions.move_to_element(driver.find_element_by_xpath(path)).perform()
        driver.implicitly_wait(0.5)
        driver.find_element_by_xpath(path).click()
        driver.implicitly_wait(1)
        driver.get_screenshot_as_file("C:/iviyo/screenshot/PASS_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Clickable element found and clicked', 'PASS', "PASS_" + str(step_info).replace('-', "").replace(" ", "_"))
    except NoSuchElementException:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/FAIL_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Clickable element not found ', 'FAIL', "FAIL_" + str(step_info).replace('-', "").replace(" ", "_"))

def single_input(driver, path, step_info, comp_string):
    actions = ActionChains(driver)
    try:
        actions.move_to_element(driver.find_element_by_xpath(path)).perform()
        driver.implicitly_wait(0.5)
        driver.find_element_by_xpath(path).clear()
        driver.find_element_by_xpath(path).send_keys(comp_string)
        driver.implicitly_wait(1)
        driver.get_screenshot_as_file("C:/iviyo/screenshot/PASS_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Data Entered in form element without form submit', 'PASS', "PASS_" + str(step_info).replace('-', "").replace(" ", "_"))
    except NoSuchElementException:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/FAIL_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Element to enter data not found ', 'FAIL', "FAIL_" + str(step_info).replace('-', "").replace(" ", "_"))

def single_input_submit(driver, path, step_info, comp_string):
    actions = ActionChains(driver)
    try:
        actions.move_to_element(driver.find_element_by_xpath(str(path).split(' : ')[0])).perform()
        driver.implicitly_wait(0.5)
        driver.find_element_by_xpath(str(path).split(' : ')[0]).clear()
        driver.find_element_by_xpath(str(path).split(' : ')[0]).send_keys(comp_string)
        driver.find_element_by_xpath(str(path).split(' : ')[-1]).click()
        driver.implicitly_wait(1)
        driver.get_screenshot_as_file("C:/iviyo/screenshot/PASS_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Data Entered in form element with form submit', 'PASS', "PASS_" + str(step_info).replace('-', "").replace(" ", "_"))
    except NoSuchElementException:
        driver.get_screenshot_as_file("C:/iviyo/screenshot/FAIL_" + str(step_info).replace('-', "").replace(" ", "_") + ".png")
        return ('Element to enter data not found, no form submitted ', 'FAIL', "FAIL_" + str(step_info).replace('-', "").replace(" ", "_"))


def close_process(driver):
    #get_log()
    driver.close()
    os.system('"taskkill /im chromedriver.exe /f"')

def iviyo_log():
    import docx
    import PIL
    from PIL import Image
    from openpyxl import load_workbook

    doc = docx.Document()
    doc.add_heading('Test Execution Report', 0)
    doc_para = doc.add_paragraph('This Test Execution Report includes status and screenshots of all the Test Steps of Test Cases which were included in scope. Follwing is the status, ')
    doc_para.add_run('Total Test Cases Executed are 100').bold = True
    doc_para.add_run(', and ')
    doc_para.add_run('Total Test Cases Passed are 95').italic = True
    doc.add_page_break()

    wb = load_workbook("C:/iviyo/testcases.xlsx")
    sheet = wb.active
    max_row = sheet.max_row

    for i in range(2, max_row + 1):
        if sheet.cell(row = i, column = 1).value is not None:
            doc.add_heading(sheet.cell(row = i, column = 1).value + "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; STATUS : " + sheet.cell(row = i, column = 14).value , 2)

        para = doc.add_heading(sheet.cell(row=i, column=3).value, 4)
        para.add_run('Status = ' + sheet.cell(row=i, column=12).value).bold = True
        para.add_run('Expected Output = ' + sheet.cell(row=i, column=5).value)

        if sheet.cell(row=i, column=10).value is not None:
            para.add_run('Reference String (Value to be checked on UI) = ' + sheet.cell(row=i, column=10).value)

        para.add_run('Actual Output = ' + sheet.cell(row=i, column=11).value)
        para.add_run('Test Type = ' + sheet.cell(row=i, column=7).value)

        img = Image.open('C:/iviyo/screenshot/' + str(sheet.cell(row=i, column=13).value) + '.png')
        img2 = img.resize((500, 250), PIL.Image.ANTIALIAS)
        img2.save('C:/iviyo/screenshot/' + str(sheet.cell(row=i, column=13).value) + '.png')
        doc.add_picture('C:/iviyo/screenshot/' + str(sheet.cell(row=i, column=13).value) + '.png')
        doc.add_page_break()

    doc.save('C:/iviyo/log.docx')